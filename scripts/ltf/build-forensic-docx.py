#!/usr/bin/env python3
"""
Builds LTF_Forensic_Report.docx — narrative forensic report for L&T Finance (FY26).
Reads: scripts/ltf/extracted/anchor_facts.json + scripts/ltf/forensic_findings.json
Writes: ~/Downloads/L&T Finance Docs/LTF_Forensic_Report.docx
Run: python3 scripts/ltf/build-forensic-docx.py
"""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
EXTR = os.path.join(HERE, "extracted")
A = json.load(open(os.path.join(EXTR, "anchor_facts.json")))
F = json.load(open(os.path.join(HERE, "forensic_findings.json")))
AR = json.load(open(os.path.join(EXTR, "annual_report_facts.json")))
OUT = os.path.expanduser("~/Downloads/L&T Finance Docs/LTF_Forensic_Report.docx")

# palette
NAVY = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x13, 0x31, 0x5C)
RED = RGBColor(0xC0, 0x39, 0x2B)
AMBER = RGBColor(0xB9, 0x6A, 0x00)
GREEN = RGBColor(0x1E, 0x84, 0x49)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
INK = RGBColor(0x1A, 0x1A, 0x1A)
SEV_HEX = {"critical": "C0392B", "high": "E67E22", "medium": "F1C40F", "low": "8DA9C4", "info": "1E8449"}
VERD_HEX = {"REFUTED": "C0392B", "PARTIALLY-CONFIRMED": "E67E22", "CONFIRMED": "1E8449", "UNVERIFIABLE": "7F8C8D"}

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)

def set_cell_text(cell, text, *, bold=False, color=None, size=9.5, white=False, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if white: run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color: run.font.color.rgb = color

def heading(text, size=15, color=NAVY, space_before=14):
    p = doc.add_paragraph()
    p.space_before = Pt(space_before)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def subheading(text):
    return heading(text, size=12, color=BLUE, space_before=10)

def body(text, italic=False, color=INK, size=10.5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.font.color.rgb = color; r.font.size = Pt(size); r.bold = bold
    return p

def callout(text, bar_color="C0392B", fill="EEF4FA"):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    shade(c, fill)
    set_cell_text(c, text, size=10.5, color=INK)
    # left accent border
    tcPr = c._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left"); left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "36"); left.set(qn("w:color"), bar_color)
    borders.append(left); tcPr.append(borders)
    return t

def style_header_row(row, hexc="13315C"):
    for cell in row.cells:
        shade(cell, hexc)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.bold = True; r.font.size = Pt(9)

# ---------------- TITLE ----------------
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = tp.add_run(F["_meta"]["title"]); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
sp = doc.add_paragraph(); r = sp.add_run(F["_meta"]["subtitle"]); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
mp = doc.add_paragraph(); r = mp.add_run(f"CMP Rs {F['valuation']['cmp']}  ·  P/B {F['valuation']['pb']:.2f}x  ·  RoE {F['valuation']['roe_fy26']:.1f}%  ·  RoA {F['valuation']['roa_fy26']:.2f}%  ·  GS3 2.88% / PCR ~67%  ·  Auditor: Unmodified")
r.font.size = Pt(10); r.font.color.rgb = BLUE; r.bold = True
sc = doc.add_paragraph(); r = sc.add_run("Red-team scorecard: " + F["_meta"]["redteam_scorecard"]); r.font.size = Pt(10); r.font.color.rgb = RED; r.bold = True

# ---------------- VERDICT ----------------
heading("1.  Forensic Verdict")
callout(F["verdict_headline"], bar_color="C0392B")
doc.add_paragraph()
subheading("Investment stance")
callout(F["stance"], bar_color="B96A00", fill="FBF3E7")
doc.add_paragraph()
callout(F["governance_verdict"], bar_color="1E8449", fill="EAF5EE")

# ---------------- RED-FLAG REGISTER ----------------
heading("2.  Red-Flag Register")
body("Severity-ranked forensic findings. Every item is traced to a primary document (audited results pack, earnings transcript, or sell-side note).", italic=True, color=GREY, size=9.5)
order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
flags = sorted(F["red_flags"], key=lambda x: order.get(x["severity"], 9))
tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, h in enumerate(["#", "Finding", "Severity", "Detail / evidence"]):
    set_cell_text(hdr[i], h, bold=True, white=True, size=9)
style_header_row(tbl.rows[0])
for f in flags:
    row = tbl.add_row().cells
    set_cell_text(row[0], f["id"], bold=True, size=9)
    set_cell_text(row[1], f["title"], bold=True, size=9.5)
    set_cell_text(row[2], f["severity"].upper(), bold=True, white=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(row[2], SEV_HEX.get(f["severity"], "7F8C8D"))
    set_cell_text(row[3], f["detail"] + "\n[Evidence: " + f["evidence"] + " | " + f["source"] + "]", size=9)
    # italicise evidence line
    for p in row[3].paragraphs[1:]:
        for rr in p.runs: rr.italic = True; rr.font.color.rgb = GREY
for col, w in zip(tbl.columns, [0.4, 2.1, 0.9, 4.0]):
    for c in col.cells: c.width = Inches(w)

subheading("Offsetting positives")
ptbl = doc.add_table(rows=1, cols=3); ptbl.style = "Table Grid"
for i, h in enumerate(["#", "Positive", "Detail / evidence"]):
    set_cell_text(ptbl.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(ptbl.rows[0], "1E8449")
for p in F["positives"]:
    row = ptbl.add_row().cells
    set_cell_text(row[0], p["id"], bold=True, size=9, color=GREEN)
    set_cell_text(row[1], p["title"], bold=True, size=9.5, color=GREEN)
    set_cell_text(row[2], p["detail"] + "\n[" + p["evidence"] + " | " + p["source"] + "]", size=9)
    for pp in row[2].paragraphs[1:]:
        for rr in pp.runs: rr.italic = True; rr.font.color.rgb = GREY
for col, w in zip(ptbl.columns, [0.4, 2.2, 4.8]):
    for c in col.cells: c.width = Inches(w)

# ---------------- RESTATED P&L ----------------
doc.add_page_break()
heading("3.  Restated P&L (NBFC format) & earnings quality")
body("Screener.in presents LTF on a manufacturing template (Raw Material, Inventory Days, OPM%) — invalid for a lender. Below is the audited P&L re-cast into proper NBFC form. Rs crore.", italic=True, color=GREY, size=9.5)
c = A["consolidated_PL_INRcr"]
def g(k, y): return c[k][y]
NII = lambda y: g("interest_income", y) - g("finance_costs", y)
tni = lambda y: NII(y) + g("fees_commission", y) + g("net_gain_fair_value", y) + g("dividend_income", y) + g("other_income", y)
opex = lambda y: g("employee_benefits", y) + g("depreciation_amort", y) + g("other_expenses", y)
ppop = lambda y: tni(y) - opex(y)
cc = lambda y: g("impairment_fin_instruments", y) + g("net_loss_derecognition_amortised_cost", y)
pl_rows = [
    ("Interest income", g("interest_income", "FY26"), g("interest_income", "FY25"), ""),
    ("Less: Finance costs", -g("finance_costs", "FY26"), -g("finance_costs", "FY25"), ""),
    ("= Net Interest Income", NII("FY26"), NII("FY25"), "+14.2%, outpaced funding costs"),
    ("Fees & commission", g("fees_commission", "FY26"), g("fees_commission", "FY25"), "+14.5% — lags loan growth +25.6%"),
    ("Treasury (net FV gain)", g("net_gain_fair_value", "FY26"), g("net_gain_fair_value", "FY25"), "-65% — tailwind faded"),
    ("Dividend + other income", g("dividend_income", "FY26") + g("other_income", "FY26"), g("dividend_income", "FY25") + g("other_income", "FY25"), ""),
    ("= Total net income", tni("FY26"), tni("FY25"), ""),
    ("Operating expenses", -opex("FY26"), -opex("FY25"), "Dep +51% (gold/tech)"),
    ("= Pre-provision operating profit", ppop("FY26"), ppop("FY25"), ""),
    ("Impairment (ECL)", -g("impairment_fin_instruments", "FY26"), -g("impairment_fin_instruments", "FY25"), "FLAT vs +25.6% loans"),
    ("Net loss on derecognition (write-offs)", -g("net_loss_derecognition_amortised_cost", "FY26"), -g("net_loss_derecognition_amortised_cost", "FY25"), "DOUBLED"),
    ("= TRUE combined credit cost", -cc("FY26"), -cc("FY25"), "ROSE +11.2% (2.60% of avg loans)"),
    ("= PBT before exceptional", g("PBT_before_exceptional", "FY26"), g("PBT_before_exceptional", "FY25"), ""),
    ("Exceptional (Labour Codes)", g("exceptional_items", "FY26"), 0, "One-time charge"),
    ("Tax", -g("tax_total", "FY26"), -g("tax_total", "FY25"), "~25.4%"),
    ("= PAT (owners)", c["PAT_owners"]["FY26"], c["PAT_owners"]["FY25"], "+12.8%"),
]
pt = doc.add_table(rows=1, cols=4); pt.style = "Table Grid"
for i, h in enumerate(["Line item", "FY26", "FY25", "Forensic note"]):
    set_cell_text(pt.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(pt.rows[0])
for label, v26, v25, note in pl_rows:
    row = pt.add_row().cells
    is_sub = label.startswith("=")
    set_cell_text(row[0], label, bold=is_sub, size=9.5, color=(NAVY if is_sub else INK))
    set_cell_text(row[1], f"{v26:,.0f}", bold=is_sub, size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_text(row[2], f"{v25:,.0f}", bold=is_sub, size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
    redword = any(w in note for w in ["FLAT", "DOUBLED", "ROSE", "lags", "faded"])
    set_cell_text(row[3], note, size=8.5, italic=True, color=(RED if redword else GREY))
    if is_sub:
        for cc_ in row: shade(cc_, "EEF4FA")
for col, w in zip(pt.columns, [3.0, 1.0, 1.0, 2.4]):
    for cell in col.cells: cell.width = Inches(w)

body("")
callout("Earnings-quality read: FY26 PAT-owners +12.8% (clean PAT-before-exceptional +13.6%) is driven by NII growth and a flat impairment line — but the flat impairment hides a doubled Rs562cr write-off/derecognition loss, so true combined credit cost rose ~11%. Layered on top: a ~Rs301cr Stage-3 overlay release and the elimination of the entire Rs125cr macro-prudential buffer in a Q4 ECL refresh. Treasury and tax were the only clean legs (both modest headwinds). Verdict: earnings quality is LOW at the margin — the growth leans on provisioning levers, not pure operating leverage.", bar_color="C0392B")

# ---------------- DuPont ----------------
heading("4.  DuPont — returns decomposition")
bs = A["consolidated_balance_sheet_INRcr"]
avgA = (bs["total_assets"]["FY26"] + bs["total_assets"]["FY25"]) / 2
avgNW = (bs["networth_consol"]["FY26"] + bs["networth_consol"]["FY25"]) / 2
roa = c["PAT_owners"]["FY26"] / avgA
roe = c["PAT_owners"]["FY26"] / avgNW
lev = avgA / avgNW
body(f"RoA (on average total assets) = {roa*100:.2f}%   ×   Leverage = {lev:.2f}x   =   RoE {roe*100:.2f}%.", bold=True, color=NAVY)
body(f"Management reports a higher 2.37% RoA (down 7bps YoY) on a friendlier denominator; both readings sit ~40-50bps below the 2.8% near-term target. At an ~11% RoE versus a ~13-14% cost of equity, the franchise is value-neutral-to-destructive at current returns — the core of the valuation problem (Section 7).", size=10)

# ---------------- ASSET QUALITY ----------------
doc.add_page_break()
heading("5.  Asset quality & provisioning — the forensic core")
callout("The headline asset-quality 'improvement' is manufactured by coverage release, write-offs and denominator growth — not credit cure, and this is now confirmed in the AUDITED annual-report ECL notes. The entire Rs575cr MFI macro-prudential buffer was unwound to NIL, ~Rs960cr of provisions on existing assets were released, and total ECL allowance FELL (Rs3,537cr → Rs3,458cr) even though gross loans grew +25.6% — cutting Stage-3 PCR from 73.0% to 68.1%. Simultaneously Rs2,413cr of loans were charged off, holding reported GS3 at 2.88% while absolute net Stage-3 actually rose +29%. ECL (incl. the overlay) is a formal Key Audit Matter; the opinion is unmodified.", bar_color="C0392B")
aq = [
    ("Gross Stage 3 (standalone)", "2.88%", "3.29% (Q4FY25)"),
    ("Net Stage 3 (abs. +29%)", "0.96%", "0.92% (QoQ ↑)"),
    ("Stage-3 PCR (audited)", "68.1%", "72.97% (cut)"),
    ("Total ECL allowance (audited)", "Rs 3,458cr / 2.89%", "Rs 3,537cr / 3.72% (fell)"),
    ("Macro-prudential buffer (audited)", "Rs NIL", "Rs 575cr (unwound)"),
    ("Loans written off (audited)", "Rs 2,413cr", "Rs 2,382cr (~2% charge-off)"),
    ("Stage-1 PCR (96% performing)", "0.80%", "0.52% (raised — forward buffer)"),
    ("Impairment (P&L)", "Rs 2,184cr", "Rs 2,193cr (flat)"),
    ("Net derecognition loss (write-offs)", "Rs 562cr", "Rs 275cr (doubled)"),
    ("TRUE combined credit cost", "Rs 2,746cr / 2.60%", "Rs 2,468cr / +11.2%"),
    ("Unsecured retail loans (audited)", "Rs 62,013cr (~53%)", "Rs 51,579cr (+20%)"),
    ("Contingent liabilities (audited)", "Rs 308cr", "Rs 326cr (small/clean)"),
    ("SR / wholesale tail (now FVTPL)", "~Rs 7,000cr", "3-4 yr drag; excl. from RoA target"),
]
at = doc.add_table(rows=1, cols=3); at.style = "Table Grid"
for i, h in enumerate(["Metric", "FY26 / latest", "Prior / note"]):
    set_cell_text(at.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(at.rows[0], "B96A00")
for m, v, p in aq:
    row = at.add_row().cells
    set_cell_text(row[0], m, bold=True, size=9.5)
    set_cell_text(row[1], v, bold=True, size=9.5, color=NAVY)
    set_cell_text(row[2], p, size=9, color=GREY)
for col, w in zip(at.columns, [2.8, 1.8, 2.8]):
    for cell in col.cells: cell.width = Inches(w)

# --- audited ECL Stage movement (annual report Note 45) ---
subheading("Audited ECL Stage 1/2/3 movement (Annual Report Note 45) — the gap now closed")
ec = AR["ecl_loans_consolidated"]
gc, al = ec["gross_carrying_by_stage"], ec["ecl_allowance_by_stage"]
et = doc.add_table(rows=1, cols=5); et.style = "Table Grid"
for i, h in enumerate(["As at 31-Mar (Rs cr)", "Stage 1", "Stage 2", "Stage 3", "Total"]):
    set_cell_text(et.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(et.rows[0], "B96A00")
def ecl_row(label, obj, bold=False):
    row = et.add_row().cells
    set_cell_text(row[0], label, bold=bold, size=9)
    for j, k in enumerate(["stage1", "stage2", "stage3", "total"]):
        set_cell_text(row[j+1], f"{obj[k]:,.0f}", size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
ecl_row("Gross carrying — FY26", gc["FY26"], True)
ecl_row("Gross carrying — FY25", gc["FY25"])
ecl_row("ECL allowance — FY26", al["FY26"], True)
ecl_row("ECL allowance — FY25", al["FY25"])
pcr26 = al["FY26"]["stage3"]/gc["FY26"]["stage3"]*100
pcr25 = al["FY25"]["stage3"]/gc["FY25"]["stage3"]*100
prow = et.add_row().cells
set_cell_text(prow[0], "Stage-3 PCR (ECL/gross)", bold=True, color=NAVY, size=9)
set_cell_text(prow[3], f"{pcr26:.1f}%", bold=True, color=RED, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell_text(prow[4], f"(FY25 {pcr25:.1f}%)", italic=True, color=RED, size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
for col, w in zip(et.columns, [2.6, 1.2, 1.2, 1.2, 1.2]):
    for cell in col.cells: cell.width = Inches(w)
body(f"Forensic levers (audited, FY26): macro-prudential provision Rs{ec['macro_prudential_provision']['FY25']}cr → NIL · provisions on existing assets released ~Rs{abs(ec['provision_release_on_existing_assets_net_of_recovery_FY26']):.0f}cr · loans written off Rs{ec['write_off_FY26']:,.0f}cr · total ECL allowance Rs{al['FY25']['total']:,.0f} → Rs{al['FY26']['total']:,.0f}cr (coverage {ec['derived']['total_ECL_coverage_FY25_pct']}% → {ec['derived']['total_ECL_coverage_FY26_pct']}%) · gold acquisition added Rs{ec['business_combination_gold_added_to_stage1_gross']:,.0f}cr to Stage-1.", size=9.5, color=RED)
body(f"Auditor & disclosure check: ECL (incl. the management overlay) is a formal Key Audit Matter ({AR['kams']['page']}); joint auditors tested PD/LGD, staging and macro overlays and issued an UNMODIFIED opinion. Related-party with parent L&T is modest and transparent — the Rs{AR['related_party_LandT']['brand_license_fee_to_LT']['FY26']}cr brand-license fee is the main item (a cost of the AAA umbrella, not abusive RPT). Contingent liabilities are small (~Rs{AR['contingent_liabilities_consolidated']['contingent_liabilities']['FY26']:.0f}cr). The provisioning choices are aggressive but were made within an audited, scrutinised framework.", size=9.5, italic=True, color=GREY)

# ---------------- MGMT vs AUDITED ----------------
heading("6.  Management headline vs verified reality")
mt = doc.add_table(rows=1, cols=4); mt.style = "Table Grid"
for i, h in enumerate(["Metric", "Management headline", "Forensic / audited reality", "Read"]):
    set_cell_text(mt.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(mt.rows[0])
for m in F["mgmt_vs_audited"]:
    row = mt.add_row().cells
    set_cell_text(row[0], m["metric"], bold=True, size=9)
    set_cell_text(row[1], m["headline"], size=9, color=GREY)
    set_cell_text(row[2], m["verified"], size=9, color=RED)
    set_cell_text(row[3], m["read"], size=9, italic=True)
for col, w in zip(mt.columns, [1.6, 2.0, 2.6, 1.8]):
    for cell in col.cells: cell.width = Inches(w)

# ---------------- SELL-SIDE & VALUATION ----------------
heading("7.  Sell-side triangulation & valuation")
sst = doc.add_table(rows=1, cols=4); sst.style = "Table Grid"
for i, h in enumerate(["House / rating", "TP & basis", "RoA path", "Note"]):
    set_cell_text(sst.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(sst.rows[0])
for s in F["sellside"]:
    row = sst.add_row().cells
    set_cell_text(row[0], f"{s['house']} — {s['rating']}", bold=True, size=9, color=(GREEN if s["rating"] == "BUY" else AMBER))
    set_cell_text(row[1], f"Rs {s['tp']} (+{s['upside_pct']}%); {s['target_basis']}", size=9)
    set_cell_text(row[2], s["roa_path"], size=9)
    set_cell_text(row[3], f"{s['roe_path']}. {s['note']}", size=8.5, color=GREY)
for col, w in zip(sst.columns, [1.6, 2.2, 2.0, 2.2]):
    for cell in col.cells: cell.width = Inches(w)
body("")
callout("Valuation verdict — " + F["valuation"]["gordon_note"] + "  " + F["valuation"]["read"], bar_color="C0392B")

# ---------------- RED-TEAM ----------------
doc.add_page_break()
heading("8.  Adversarial red-team — bull pillars stress-tested")
body("Each pillar was attacked by 3 independent skeptics using only the primary documents. " + F["_meta"]["redteam_scorecard"], italic=True, color=GREY, size=9.5)
rt = doc.add_table(rows=1, cols=3); rt.style = "Table Grid"
for i, h in enumerate(["Bull pillar", "Verdict", "Why it fails / strongest refutation"]):
    set_cell_text(rt.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(rt.rows[0])
for p in F["pillars"]:
    row = rt.add_row().cells
    set_cell_text(row[0], p["pillar"], bold=True, size=9)
    set_cell_text(row[1], f"{p['verdict']}\n({p['votes']})", bold=True, white=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(row[1], VERD_HEX.get(p["verdict"], "7F8C8D"))
    set_cell_text(row[2], p["summary"] + "\n→ " + p["refutation"], size=9)
    for pp in row[2].paragraphs[1:]:
        for rr in pp.runs: rr.font.color.rgb = RED; rr.font.size = Pt(8.5)
for col, w in zip(rt.columns, [2.4, 1.4, 4.0]):
    for cell in col.cells: cell.width = Inches(w)

# ---------------- GOVERNANCE / CONTROLS / REGULATORY ----------------
heading("9.  Governance, controls & regulatory standing (CARO · IFC · RBI)")
body("The integrity counterweight: the issues in this report are about earnings quality, not compliance. CARO, IFC, the RBI inspection, penalties, ratings and fraud disclosures are clean.", italic=True, color=GREY, size=9.5)
gt = doc.add_table(rows=1, cols=4); gt.style = "Table Grid"
for i, h in enumerate(["Item", "Status", "Detail", "Source"]):
    set_cell_text(gt.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(gt.rows[0], "1E8449")
GSTAT = {"CLEAN": "1E8449", "STRONG": "1E8449", "MINOR": "8DA9C4", "WATCH": "E67E22", "ADVERSE": "C0392B"}
for g in F["governance_controls"]:
    row = gt.add_row().cells
    set_cell_text(row[0], g["item"], bold=True, size=9)
    set_cell_text(row[1], g["status"], bold=True, white=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(row[1], GSTAT.get(g["status"], "7F8C8D"))
    set_cell_text(row[2], g["detail"], size=9)
    set_cell_text(row[3], g["source"], size=8.5, italic=True, color=GREY)
for col, w in zip(gt.columns, [1.9, 0.9, 4.0, 1.3]):
    for cell in col.cells: cell.width = Inches(w)

# ---------------- DATA INTEGRITY ----------------
heading("10.  Data-integrity log")
for item in F["data_integrity_log"]:
    p = doc.add_paragraph(style="List Bullet"); r = p.add_run(item); r.font.size = Pt(9.5)

# ---------------- MONITORABLES ----------------
heading("11.  Key monitorables")
for m in F["monitorables"]:
    p = doc.add_paragraph(style="List Bullet"); r = p.add_run(m); r.font.size = Pt(10)

# ---------------- SOURCES ----------------
heading("12.  Source index")
stbl = doc.add_table(rows=1, cols=3); stbl.style = "Table Grid"
for i, h in enumerate(["Document", "Role", "Type"]):
    set_cell_text(stbl.rows[0].cells[i], h, bold=True, white=True, size=9)
style_header_row(stbl.rows[0], "7F8C8D")
for s in F["sources"]:
    row = stbl.add_row().cells
    set_cell_text(row[0], s["doc"], size=9)
    set_cell_text(row[1], s["role"], size=9)
    set_cell_text(row[2], s["type"], size=9)
for col, w in zip(stbl.columns, [3.4, 3.0, 1.4]):
    for cell in col.cells: cell.width = Inches(w)

disc = doc.add_paragraph()
r = disc.add_run("Prepared from primary documents for internal research use. Audited figures cross-footed and independently re-parsed (qwen3:14b) to the rupee. Not investment advice. Figures in Rs crore unless stated.")
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY

doc.save(OUT)
print("WROTE:", OUT)
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))
